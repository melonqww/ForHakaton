"""
doc_generator.py
Генерация аналитических документов (DOCX + PDF) и executive summary для чиновников.
Использует Qwen 1.5B через Ollama для генерации текстов.
"""

from __future__ import annotations

import io
import os
import time
import datetime
import requests


# ── Ollama ──────────────────────────────────────────────────────────────────

_OLLAMA_MODEL_LARGE = "qwen2.5:1.5b"


def _call_qwen(system_prompt: str, user_prompt: str, ollama_url: str, timeout: int = 90) -> str:
    """Вызов Qwen 1.5B через Ollama API/Chat. Возвращает пустую строку при ошибке."""
    try:
        # Enforce Russian language strictly to avoid Chinese/gibberish drift
        strict_system = system_prompt + "\nОТВЕЧАЙ СТРОГО НА РУССКОМ ЯЗЫКЕ! Использование китайских иероглифов, латиницы или других нерусских символов категорически запрещено."
        chat_url = ollama_url.replace("/api/generate", "/api/chat")
        resp = requests.post(
            chat_url,
            json={
                "model": _OLLAMA_MODEL_LARGE,
                "messages": [
                    {"role": "system", "content": strict_system},
                    {"role": "user", "content": user_prompt}
                ],
                "stream": False,
                "options": {"temperature": 0.3, "repetition_penalty": 1.05, "num_predict": 2048},
            },
            timeout=timeout,
        )
        if resp.status_code == 200:
            text = resp.json().get("message", {}).get("content", "").strip()
            text = text.replace('"', '').replace("'", "")
            return text
    except Exception:
        pass
    return ""


# ── Executive Summary (для вкладки Итоги) ───────────────────────────────────

def generate_executive_summary(stats: dict, ollama_url: str) -> str:
    """
    Краткая управленческая сводка ~400-600 символов.
    Вызывается автоматически после обработки файла.
    """
    total = stats.get("total_count", 0)
    problems = stats.get("problems_count", 0)
    top3 = stats.get("top3_districts", [])
    cats = stats.get("category_counts", {})
    ranks = stats.get("rank_counts", {})

    # Считаем критические (ранг 4-5)
    critical = sum(v for k, v in ranks.items() if int(k) >= 4)

    top3_str = ", ".join(
        f"{d['district']} ({d['count']} обр.)" for d in top3[:3]
    ) or "нет данных"

    top_cats = sorted(cats.items(), key=lambda x: -x[1])[:5]
    cats_str = ", ".join(f"{k}: {v}" for k, v in top_cats) or "нет данных"

    system_prompt = (
        "Ты — главный аналитик Центра управления регионом Омской области. Напиши структурированную, "
        "краткую и легкую для чтения аналитическую сводку для губернатора региона по итогам анализа обращений граждан. "
        "Раздели ответ на три небольших абзаца (каждый по 1-2 предложения, разделяй их пустой строкой):\n"
        "1. Общая оценка ситуации: общая характеристика и масштаб проблем («Что и как в целом»).\n"
        "2. Зоны риска: кратко выдели самые критические сферы и наиболее проблемные районы.\n"
        "3. Предлагаемые решения: конкретные управленческие шаги по снижению социальной напряженности.\n\n"
        "Пиши строго профессиональным языком, не придумывай новые факты. Текст должен легко читаться благодаря абзацам."
    )
    user_prompt = (
        f"Данные для анализа:\n"
        f"- Всего обращений: {total}, из них реальные проблемы: {problems}, "
        f"критические (ранг 4-5): {critical}\n"
        f"- Наиболее проблемные районы (ТОП-3): {top3_str}\n"
        f"- Главные темы жалоб: {cats_str}\n\n"
        f"Напиши структурированную сводку для губернатора по указанному формату:"
    )

    text = _call_qwen(system_prompt, user_prompt, ollama_url, timeout=60)

    if not text:
        # Fallback без ИИ
        pct = f"{problems / total * 100:.1f}%" if total else "0%"
        top_name = top3[0]["district"] if top3 else "—"
        top_cat = top_cats[0][0] if top_cats else "—"
        return (
            f"За анализируемый период поступило {total:,} обращений, "
            f"из которых {problems:,} ({pct}) квалифицированы как реальные проблемы, "
            f"требующие решения. "
            f"Критических обращений (ранг 4–5): {critical:,}. "
            f"Наибольшее число обращений зафиксировано в районе «{top_name}». "
            f"Доминирующая тема: «{top_cat}»."
        )

    return text[:1200]


# ── DOCX Generator ───────────────────────────────────────────────────────────

def generate_docx(
    stats: dict,
    ai_summary: str,
    source_file_name: str,
    ollama_url: str,
) -> bytes:
    """
    Генерирует аналитический Word-документ для чиновников.
    Возвращает bytes (содержимое .docx файла).
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return b""

    doc = Document()

    # Настройка страницы
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    def _set_style(run, bold=False, size=11, color=None):
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _heading(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        _set_style(run, bold=True, size=14 if level == 1 else 12,
                   color=(30, 58, 138) if level == 1 else (55, 65, 81))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        return p

    def _para(text, size=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(4)
        return p

    def _add_formatted_text(p, line_text):
        parts = line_text.split("**")
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            run.font.size = Pt(11)
            if idx % 2 == 1:
                run.bold = True

    def _add_markdown(text):
        if not text:
            return
        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 1. Заголовки (### или ## или #)
            if line.startswith("###") or line.startswith("##") or line.startswith("#"):
                clean_line = line.lstrip("#").strip()
                clean_line = clean_line.replace("**", "").replace("*", "")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(clean_line)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(30, 58, 138)
            
            # 2. Списки (- или * или 1.)
            elif line.startswith("- ") or line.startswith("* ") or (line[0:1].isdigit() and ". " in line[:4]):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(3)
                
                parts = line.split(" ", 1)
                marker = parts[0]
                content = parts[1] if len(parts) > 1 else ""
                
                run_marker = p.add_run(marker + " ")
                run_marker.bold = True
                run_marker.font.size = Pt(11)
                
                _add_formatted_text(p, content)
                
            # 3. Обычный абзац
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(4)
                _add_formatted_text(p, line)


    def _add_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        # Заголовок
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            # Синий фон
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1E3A8A')
            tcPr.append(shd)
        # Данные
        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx + 1]
            for c_idx, val in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = str(val)
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                # Чередование строк
                if r_idx % 2 == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'EFF6FF')
                    tcPr.append(shd)
        return table

    # ── Шапка документа ────────────────────────────────────────────────────
    now = datetime.datetime.now()
    date_str = now.strftime("%d.%m.%Y")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("АНАЛИТИЧЕСКАЯ СПРАВКА")
    _set_style(run, bold=True, size=18, color=(30, 58, 138))

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "Район:" in source_file_name:
        run2 = p_sub.add_run(f"по обращениям граждан Омской области — {source_file_name}")
    else:
        run2 = p_sub.add_run("по обращениям граждан Омской области")
    _set_style(run2, size=13, color=(75, 85, 99))

    doc.add_paragraph()

    # Мета-инфо
    meta_tbl = doc.add_table(rows=3, cols=2)
    meta_tbl.style = "Table Grid"
    meta_rows = [
        ("Дата формирования:", date_str),
        ("Источник данных:", os.path.basename(source_file_name)),
        ("Сформировано системой:", "Анализ обращений граждан — ЦУР Омской области"),
    ]
    for i, (label, val) in enumerate(meta_rows):
        meta_tbl.rows[i].cells[0].text = label
        meta_tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        meta_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        meta_tbl.rows[i].cells[1].text = val
        meta_tbl.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── 1. Управленческая сводка ────────────────────────────────────────────
    _heading("1. Управленческая сводка", level=1)
    if ai_summary:
        _add_markdown(ai_summary)
    else:
        _para("Данные аналитики недоступны. См. статистику ниже.", italic=True)

    # ── 2. Общая статистика ─────────────────────────────────────────────────
    _heading("2. Общая статистика", level=1)

    total = stats.get("total_count", 0)
    problems = stats.get("problems_count", 0)
    spam = total - problems
    ranks = stats.get("rank_counts", {})
    critical = sum(v for k, v in ranks.items() if int(k) >= 4)
    pct = f"{problems / total * 100:.1f}%" if total else "0%"

    _add_table(
        ["Показатель", "Значение"],
        [
            ["Всего обращений", f"{total:,}"],
            ["Реальные проблемы", f"{problems:,} ({pct})"],
            ["Не требует решения", f"{spam:,}"],
            ["Критические (ранг 4–5)", f"{critical:,}"],
        ],
    )

    # ── 3. ТОП-3 проблемных района ──────────────────────────────────────────
    _heading("3. Наиболее проблемные районы", level=1)
    top10 = stats.get("top10_districts", [])
    if top10:
        rows_d = []
        medals = ["🥇", "🥈", "🥉"]
        for i, d in enumerate(top10[:10]):
            avg_r = d.get("avg_rank", 0)
            rows_d.append([
                f"{medals[i]} {d['district']}" if i < 3 else d["district"],
                str(d["count"]),
                f"{avg_r:.1f}",
                str(d.get("critical_count", 0)),
            ])
        _add_table(
            ["Район", "Обращений", "Ср. ранг", "Критических"],
            rows_d,
        )

    # Нарратив по топ-3 от Qwen
    top3 = stats.get("top3_districts", [])
    if top3:
        doc.add_paragraph()
        _heading("Аналитика по ключевым районам", level=2)

        top3_facts = "\n".join(
            f"- {d['district']}: {d['count']} обращений, "
            f"критических: {d.get('critical_count', 0)}, "
            f"темы: {'; '.join(list(d.get('categories', {}).keys())[:3])}"
            for d in top3
        )
        system_prompt = (
            "Ты — аналитик Центра управления регионом Омской области. "
            "Напиши один профессиональный аналитический абзац о трех самых проблемных районах региона для губернатора. "
            "Пиши на русском языке, без заголовков, без списков, сплошным текстом (5-6 предложений). "
            "Объясни причины большого числа жалоб, какие категории доминируют и что требует вмешательства."
        )
        user_prompt = f"Данные по районам:\n{top3_facts}"
        narrative = _call_qwen(system_prompt, user_prompt, ollama_url, timeout=90)
        if narrative:
            _add_markdown(narrative)

    # ── 4. Распределение по темам ────────────────────────────────────────────
    _heading("4. Тематическое распределение обращений", level=1)
    cats = stats.get("category_counts", {})
    if cats:
        top_cats = sorted(cats.items(), key=lambda x: -x[1])[:15]
        _add_table(
            ["Тема", "Количество", "Доля"],
            [
                [cat, f"{cnt:,}", f"{cnt / problems * 100:.1f}%" if problems else "—"]
                for cat, cnt in top_cats
            ],
        )

    # ── 5. Рекомендации от ИИ ───────────────────────────────────────────────
    _heading("5. Рекомендации", level=1)

    if cats and top3:
        top_cat_names = [c[0] for c in sorted(cats.items(), key=lambda x: -x[1])[:5]]
        system_prompt = (
            "Ты — старший аналитик Центра управления регионом Омской области. "
            "На основе данных сформулируй 4-5 конкретных управленческих рекомендаций для руководства региона. "
            "Каждая рекомендация должна быть одним конкретным действием на русском языке, без вводных фраз и общих слов (списком)."
        )
        user_prompt = (
            f"Данные для анализа:\n"
            f"- Всего проблем: {problems}\n"
            f"- Из них критических: {critical}\n"
            f"- Главные темы жалоб: {', '.join(top_cat_names)}\n"
            f"- Наиболее проблемный район: {top3[0]['district'] if top3 else '—'}"
        )
        recommendations = _call_qwen(system_prompt, user_prompt, ollama_url, timeout=90)
        if recommendations:
            _add_markdown(recommendations)
        else:
            _add_markdown(
                "1. Усилить контроль за устранением критических обращений (ранг 4–5) "
                "в районах с наибольшим числом жалоб.\n"
                "2. Провести выездные проверки в ТОП-3 проблемных районах.\n"
                "3. Организовать рабочее совещание с профильными ведомствами "
                "по доминирующим темам обращений."
            )

    # ── Подпись ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    _para(f"Документ сформирован автоматически {date_str}.", italic=True,
          size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para("Система анализа обращений граждан — ЦУР Омской области",
          italic=True, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF Generator ────────────────────────────────────────────────────────────

def generate_pdf(
    stats: dict,
    ai_summary: str,
    source_file_name: str,
    ollama_url: str,
) -> bytes:
    """
    Генерирует PDF-версию аналитической справки через reportlab.
    Не требует MS Word.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return b""

    # Попытка зарегистрировать системный шрифт с кириллицей
    _font_name = "Helvetica"
    _font_bold = "Helvetica-Bold"
    _font_italic = "Helvetica-Oblique"
    try:
        import os as _os
        font_candidates = [
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\times.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        ]
        font_bold_candidates = [
            r"C:\Windows\Fonts\arialbd.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        ]
        for fp in font_candidates:
            if _os.path.exists(fp):
                pdfmetrics.registerFont(TTFont("CyrFont", fp))
                _font_name = "CyrFont"
                break
        for fp in font_bold_candidates:
            if _os.path.exists(fp):
                pdfmetrics.registerFont(TTFont("CyrFontB", fp))
                _font_bold = "CyrFontB"
                break
        if _font_bold == "Helvetica-Bold" and _font_name != "Helvetica":
            _font_bold = _font_name
    except Exception:
        pass

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    W = A4[0] - 4.5 * cm  # ширина контента

    ST_TITLE = ParagraphStyle("title", fontName=_font_bold, fontSize=18,
                               textColor=colors.HexColor("#1E3A8A"),
                               alignment=TA_CENTER, spaceAfter=4)
    ST_SUB = ParagraphStyle("sub", fontName=_font_name, fontSize=13,
                             textColor=colors.HexColor("#4B5563"),
                             alignment=TA_CENTER, spaceAfter=16)
    ST_H1 = ParagraphStyle("h1", fontName=_font_bold, fontSize=13,
                            textColor=colors.HexColor("#1E3A8A"),
                            spaceBefore=14, spaceAfter=6)
    ST_H2 = ParagraphStyle("h2", fontName=_font_bold, fontSize=11,
                            textColor=colors.HexColor("#374151"),
                            spaceBefore=8, spaceAfter=4)
    ST_BODY = ParagraphStyle("body", fontName=_font_name, fontSize=10,
                              leading=15, alignment=TA_JUSTIFY, spaceAfter=6)
    ST_SMALL = ParagraphStyle("small", fontName=_font_name, fontSize=8,
                               textColor=colors.grey, alignment=TA_RIGHT)
    ST_META_LBL = ParagraphStyle("ml", fontName=_font_bold, fontSize=9)
    ST_META_VAL = ParagraphStyle("mv", fontName=_font_name, fontSize=9)
    ST_TBL_HDR = ParagraphStyle("tbl_hdr", fontName=_font_bold, fontSize=9,
                                 textColor=colors.white, alignment=TA_LEFT)
    ST_TBL_CELL = ParagraphStyle("tbl_cell", fontName=_font_name, fontSize=9,
                                  textColor=colors.HexColor("#1A202C"), alignment=TA_LEFT)
    ST_LIST = ParagraphStyle("list", fontName=_font_name, fontSize=10,
                              leading=14, alignment=TA_LEFT, leftIndent=15, spaceAfter=4)

    def _markdown_to_html_helper(t):
        # Escape XML entities first
        t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        parts = t.split("**")
        html_parts = []
        for idx, part in enumerate(parts):
            if idx % 2 == 1:
                html_parts.append(f"<b>{part}</b>")
            else:
                html_parts.append(part)
        return "".join(html_parts)

    def _add_markdown_pdf(text, body_style=ST_BODY, h2_style=ST_H2, list_style=ST_LIST):
        if not text:
            return
        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 1. Заголовки (### или ## или #)
            if line.startswith("###") or line.startswith("##") or line.startswith("#"):
                clean_line = line.lstrip("#").strip()
                clean_line = clean_line.replace("**", "").replace("*", "")
                clean_line = clean_line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(clean_line, h2_style))
                story.append(Spacer(1, 2))
                
            # 2. Списки (- или * или 1.)
            elif line.startswith("- ") or line.startswith("* ") or (line[0:1].isdigit() and ". " in line[:4]):
                parts = line.split(" ", 1)
                marker = parts[0]
                content = parts[1] if len(parts) > 1 else ""
                
                formatted_content = _markdown_to_html_helper(content)
                text_html = f"<b>{marker}</b> {formatted_content}"
                
                story.append(Paragraph(text_html, list_style))
                story.append(Spacer(1, 2))
                
            # 3. Обычный абзац
            else:
                formatted_content = _markdown_to_html_helper(line)
                story.append(Paragraph(formatted_content, body_style))
                story.append(Spacer(1, 3))


    def _wrap_table_cells(data):
        wrapped = []
        for r_idx, row in enumerate(data):
            wrapped_row = []
            for cell in row:
                style = ST_TBL_HDR if r_idx == 0 else ST_TBL_CELL
                wrapped_row.append(Paragraph(str(cell), style))
            wrapped.append(wrapped_row)
        return wrapped

    tbl_header = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), _font_name),
        ("FONTNAME", (0, 0), (-1, 0), _font_bold),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EFF6FF")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ])

    now = datetime.datetime.now()
    date_str = now.strftime("%d.%m.%Y")
    total = stats.get("total_count", 0)
    problems = stats.get("problems_count", 0)
    spam = total - problems
    ranks = stats.get("rank_counts", {})
    critical = sum(v for k, v in ranks.items() if int(k) >= 4)
    pct = f"{problems / total * 100:.1f}%" if total else "0%"
    cats = stats.get("category_counts", {})
    top3 = stats.get("top3_districts", [])
    top10 = stats.get("top10_districts", [])

    story = []

    # Заголовок
    story.append(Paragraph("АНАЛИТИЧЕСКАЯ СПРАВКА", ST_TITLE))
    if "Район:" in source_file_name:
        story.append(Paragraph(f"по обращениям граждан Омской области — {source_file_name}", ST_SUB))
    else:
        story.append(Paragraph("по обращениям граждан Омской области", ST_SUB))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1E3A8A")))
    story.append(Spacer(1, 8))

    # Мета
    meta_data = [
        [Paragraph("Дата формирования:", ST_META_LBL), Paragraph(date_str, ST_META_VAL)],
        [Paragraph("Источник данных:", ST_META_LBL),
         Paragraph(os.path.basename(source_file_name), ST_META_VAL)],
        [Paragraph("Сформировано:", ST_META_LBL),
         Paragraph("Система анализа обращений граждан — ЦУР Омской области", ST_META_VAL)],
    ]
    meta_tbl = Table(meta_data, colWidths=[4 * cm, W - 4 * cm])
    meta_tbl.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(meta_tbl)
    story.append(Spacer(1, 12))

    # 1. Управленческая сводка
    story.append(Paragraph("1. Управленческая сводка", ST_H1))
    if ai_summary:
        _add_markdown_pdf(ai_summary)
    else:
        story.append(Paragraph("Данные аналитики недоступны. Проверьте доступность Ollama.", ST_BODY))

    # 2. Общая статистика
    story.append(Paragraph("2. Общая статистика", ST_H1))
    stat_data = [
        ["Показатель", "Значение"],
        ["Всего обращений", f"{total:,}"],
        ["Реальные проблемы", f"{problems:,} ({pct})"],
        ["Не требует решения", f"{spam:,}"],
        ["Критические (ранг 4–5)", f"{critical:,}"],
    ]
    st_tbl = Table(_wrap_table_cells(stat_data), colWidths=[W * 0.65, W * 0.35])
    st_tbl.setStyle(tbl_header)
    story.append(st_tbl)

    # 3. ТОП-10 районов
    story.append(Paragraph("3. Наиболее проблемные районы", ST_H1))
    if top10:
        medals = ["1.", "2.", "3."]
        dist_data = [["Район", "Обращений", "Ср. ранг", "Критических"]]
        for i, d in enumerate(top10[:10]):
            prefix = medals[i] if i < 3 else f"{i+1}."
            dist_data.append([
                f"{prefix} {d['district']}",
                f"{d['count']:,}",
                f"{d.get('avg_rank', 0):.1f}",
                f"{d.get('critical_count', 0):,}",
            ])
        d_tbl = Table(_wrap_table_cells(dist_data), colWidths=[W * 0.5, W * 0.18, W * 0.17, W * 0.15])
        d_tbl.setStyle(tbl_header)
        story.append(d_tbl)

    # Нарратив по топ-3
    if top3:
        story.append(Spacer(1, 6))
        story.append(Paragraph("Аналитика по ключевым районам", ST_H2))
        top3_facts = "\n".join(
            f"- {d['district']}: {d['count']} обращений, "
            f"критических: {d.get('critical_count', 0)}"
            for d in top3
        )
        system_prompt = (
            "Ты — аналитик Центра управления регионом Омской области. "
            "Напиши аналитический абзац о трех самых проблемных районах. Объясни причины большого числа жалоб. "
            "Пиши профессиональным русским языком, без заголовков, без списков, сплошным текстом (4-5 предложений)."
        )
        user_prompt = f"Данные по районам:\n{top3_facts}"
        narrative = _call_qwen(system_prompt, user_prompt, ollama_url, timeout=90)
        if narrative:
            _add_markdown_pdf(narrative)

    # 4. Тематика
    story.append(Paragraph("4. Тематическое распределение", ST_H1))
    if cats:
        top_cats = sorted(cats.items(), key=lambda x: -x[1])[:15]
        cats_data = [["Тема", "Количество", "Доля"]]
        for cat, cnt in top_cats:
            share = f"{cnt / problems * 100:.1f}%" if problems else "—"
            cats_data.append([cat, f"{cnt:,}", share])
        c_tbl = Table(_wrap_table_cells(cats_data), colWidths=[W * 0.6, W * 0.2, W * 0.2])
        c_tbl.setStyle(tbl_header)
        story.append(c_tbl)

    # 5. Рекомендации
    story.append(Paragraph("5. Рекомендации", ST_H1))
    if cats and top3:
        top_cat_names = [c[0] for c in sorted(cats.items(), key=lambda x: -x[1])[:5]]
        system_prompt = (
            "Ты — аналитик Центра управления регионом Омской области. "
            "Сформулируй 4-5 конкретных управленческих рекомендаций для руководства Омской области. "
            "Каждая рекомендация должна быть одним конкретным действием на русском языке, без общих слов и вводных фраз (списком)."
        )
        user_prompt = (
            f"Данные для анализа:\n"
            f"- Всего проблем: {problems}\n"
            f"- Критических: {critical}\n"
            f"- Главные темы: {', '.join(top_cat_names)}\n"
            f"- Самый проблемный район: {top3[0]['district']}"
        )
        recommendations = _call_qwen(system_prompt, user_prompt, ollama_url, timeout=90)
        if recommendations:
            _add_markdown_pdf(recommendations)
        else:
            _add_markdown_pdf(
                "1. Провести выездные проверки в ТОП-3 проблемных районах.\n"
                "2. Усилить контроль за устранением критических обращений (ранг 4–5).\n"
                "3. Провести совещание с профильными ведомствами по доминирующим темам."
            )

    # Подпись
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        f"Документ сформирован автоматически {date_str}. "
        "Система анализа обращений граждан — ЦУР Омской области.",
        ST_SMALL,
    ))

    doc.build(story)
    return buf.getvalue()
